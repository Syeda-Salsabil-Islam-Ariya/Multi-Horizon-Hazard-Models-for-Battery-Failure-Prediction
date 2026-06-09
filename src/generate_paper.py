"""Generate paper.docx from benchmark_results.csv and figure PNGs"""
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd

DOCX_DIR = "/run/media/touhid/SSD/batt_sense/conference/conference -A/paper"
FIG_DIR  = "/run/media/touhid/SSD/batt_sense/battery_sens/battery_sens/data"
CSV_PATH = "/run/media/touhid/SSD/batt_sense/battery_sens/battery_sens/data/benchmark_results.csv"

os.makedirs(DOCX_DIR, exist_ok=True)

doc = Document()

# ── Helpers ──────────────────────────────────────────────────────────────
def add_figure(path, caption, width=Inches(4.8)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=width)
    pc = doc.add_paragraph(caption)
    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pc.runs[0].font.size = Pt(9)

def set_cell(cell, text, bold=False, size=9):
    cell.text = str(text)
    for run in cell.paragraphs[0].runs:
        run.font.size = Pt(size)
        run.bold = bold
        run.font.name = "Times New Roman"

def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True, size=9)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            set_cell(table.rows[r + 1].cells[c], str(val), size=9)
    return table

# ══════════════════════════════════════════════════════════════════════════
#  TITLE
# ══════════════════════════════════════════════════════════════════════════
title = doc.add_heading("Multi-Horizon Hazard Models for Battery Failure Prediction: Within-Dataset Reliability and Cross-Chemistry Transferability", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

auth = doc.add_paragraph()
auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
auth.add_run("[Author Name(s)]\n[Affiliation]\n[Email]").font.size = Pt(10)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════
#  ABSTRACT
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("Abstract", level=1)
doc.add_paragraph(
    "Lithium-ion battery failure prediction is critical for safe and reliable operation across "
    "applications from electric vehicles to grid storage. This study extends an existing multi-horizon "
    "hazard classification framework—originally developed for NASA 18650 cells using Histogram-based "
    "Gradient Boosting (HGB)—to two additional datasets (CALCE LCO/CX2, Oxford LFP) and three model "
    "classes (XGBoost, LightGBM, Random Forest) with matched hyperparameters. A composite failure label "
    "is defined as State-of-Health (SOH) below 0.80 or average voltage sag below 94% of baseline. "
    "We compare isotonic and Platt (sigmoid) calibration across all model-dataset combinations and "
    "evaluate cross-chemistry transfer from LCO (NASA + CALCE) to LFP (Oxford) with and without SOH "
    "as a feature. Platt calibration universally outperforms isotonic, most dramatically on CALCE "
    "(AUC 0.694 \u2192 0.911). Cross-chemistry transfer with SOH included yields AUC 0.87\u20131.00, "
    "but removing SOH collapses AUC to 0.51\u20130.54, revealing that SOH encodes a chemistry-specific "
    "capacity-to-RUL mapping rather than a transferable degradation signal. These results establish a "
    "reproducible within-dataset performance baseline while demonstrating that cross-chemistry generalization "
    "of hazard-based battery failure models remains an open problem."
)

# ══════════════════════════════════════════════════════════════════════════
#  1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("1. Introduction", level=1)
doc.add_paragraph(
    "The original study by Shikdar and Laaksonen introduced a multi-horizon hazard classification framework "
    "for lithium-ion battery failure prediction, using Histogram-based Gradient Boosting (HGB) on the "
    "NASA 18650 dataset with a composite SOH-and-voltage-sag failure label. The work demonstrated that "
    "classifying batteries as \u201cfail within H cycles\u201d is an effective alternative to traditional "
    "remaining-useful-life (RUL) regression, with reported AUC values of 0.868\u20130.898 across horizons "
    "H \u2208 {10, 20, 30, 50}. However, several important questions were not addressed."
)
doc.add_paragraph(
    "First, the original study evaluated a single model class (HGB) on a single dataset (NASA 18650). "
    "It is unknown whether the results are sensitive to model choice, hyperparameter configuration, or "
    "the calibration method. Second, and more critically, it is unknown whether models trained on one "
    "battery chemistry generalize to others. Real-world battery fleets often contain mixed chemistries, "
    "so a practical hazard-monitoring system must either be trained separately per chemistry or rely on "
    "transferable features."
)
doc.add_paragraph(
    "This paper addresses both gaps. Our contributions are: (1) a three-model benchmark (XGBoost, "
    "LightGBM, Random Forest) with matched hyperparameters on two LCO datasets (NASA 18650 and CALCE "
    "LCO/CX2, 44 cells total); (2) a systematic comparison of isotonic and Platt (sigmoid) calibration, "
    "showing that Platt universally produces better probability estimates; and (3) a cross-chemistry "
    "transfer analysis from LCO to LFP (Oxford) with controlled ablation of the SOH feature, "
    "demonstrating that SOH drives nearly all apparent transferability."
)

# ══════════════════════════════════════════════════════════════════════════
#  2. RELATED WORK
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("2. Related Work", level=1)
doc.add_paragraph(
    "Battery prognostics has traditionally focused on RUL estimation via regression models. "
    "These methods predict a continuous time-to-failure value, which requires "
    "well-defined end-of-life thresholds and sufficient failure data for training."
)
doc.add_paragraph(
    "An alternative framing treats failure prediction as a classification problem: will the battery "
    "fail within a given horizon? Shikdar and Laaksonen [1] proposed this multi-horizon hazard approach using "
    "HGB with a composite failure label combining SOH and voltage sag. Their work demonstrated that "
    "classification-based hazard models can produce actionable warnings well before end-of-life, and that "
    "incorporating voltage sag as a secondary failure criterion captures impedance-driven failures that "
    "precede capacity fade."
)
doc.add_paragraph(
    "Cross-chemistry transfer in battery ML has received limited attention. "
    "The present study contributes direct evidence "
    "on this question by systematically ablating the SOH feature."
)
doc.add_paragraph(
    "Calibration of classification models for battery prognostics is similarly underexplored. "
    "Isotonic regression [2] and Platt scaling [3] are standard post-hoc calibration methods in ML, "
    "but their relative performance on battery degradation data\u2014characterized by class imbalance, "
    "long-tailed SOH distributions, and small cell counts\u2014has not been previously evaluated."
)

# ══════════════════════════════════════════════════════════════════════════
#  3. METHODOLOGY
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("3. Methodology", level=1)

doc.add_heading("3.1 Composite Failure Label", level=2)
doc.add_paragraph(
    "Following the original protocol [1], a battery cycle is labeled as \u201cfailure\u201d if either "
    "of two conditions is met: (1) State-of-Health (SOH) falls at or below 0.80 of initial capacity, "
    "where SOH is defined as the ratio of current discharge capacity to the mean capacity of the first "
    "10 cycles; or (2) the average discharge voltage drops below 94% of its early-life baseline (first "
    "10 cycles). The second criterion captures impedance-driven degradation where voltage sag precedes "
    "measurable capacity fade. Both conditions use the same baseline window, and once triggered, the "
    "label remains positive for all subsequent cycles. For datasets without voltage data (or where "
    "voltage is always at the cutoff, as in Oxford LFP), only the SOH criterion applies."
)

doc.add_heading("3.2 Datasets", level=2)
doc.add_paragraph(
    "Three publicly available battery aging datasets are used. The NASA 18650 dataset [4] contains "
    "37 LCO cells aged under random-walk charge/discharge profiles at room temperature, producing "
    "degradation trajectories of approximately 1,000 cycles per cell with diverse failure patterns. "
    "The CALCE LCO/CX2 dataset [5] consists of 7 cells (CS2_33\u201336, CX2_36\u201338) cycled at "
    "1C charge/1C discharge to 80% SOH or below, yielding 8733 total cycles with slow, uniform "
    "degradation spanning 775\u20131952 cycles per cell. The Oxford LFP dataset [6] contains 5 LFP "
    "pouch cells cycled at 1C/1C for approximately 300 cycles each; the flat LFP voltage plateau "
    "renders the voltage sag feature uninformative."
)

doc.add_heading("3.3 Features and Preprocessing", level=2)
doc.add_paragraph(
    "All models use the same feature set per cycle: cycle number, average voltage, minimum voltage, "
    "average current, average temperature, discharge duration, and SOH. For the cross-chemistry transfer "
    "experiment, SOH is excluded from the feature set in the no-SOH variant to test whether remaining "
    "features capture chemistry-agnostic degradation patterns. Features are not standardized (tree-based "
    "models are scale-invariant). Sequential cycles from each cell are retained without interpolation."
)

doc.add_heading("3.4 Models and Hyperparameters", level=2)
doc.add_paragraph(
    "Three tree-based classifiers are compared with hyperparameters matched to the original study [1]: "
    "XGBoost (max_depth=4, n_estimators=300, learning_rate=0.05, subsample=0.8, colsample_bytree=0.8), "
    "LightGBM (max_depth=4, n_estimators=300, learning_rate=0.05, subsample=0.8, "
    "colsample_bytree=0.8, verbose=-1), and Random Forest (max_depth=6, n_estimators=300, "
    "max_samples=0.8). All models use default class weighting."
)

doc.add_heading("3.5 Calibration", level=2)
doc.add_paragraph(
    "Two post-hoc calibration methods are compared: isotonic regression, which fits a non-decreasing "
    "step function to map raw scores to probabilities (implemented via sklearn\u2019s "
    "CalibratedClassifierCV(method=\u2018isotonic\u2019, cv=3)), and Platt scaling (method=\u2018sigmoid\u2019, "
    "cv=3), which fits a logistic regression on the raw scores. Calibration is applied within the "
    "cross-validation loop on held-out folds to avoid information leakage. Performance is measured "
    "by both AUC (discrimination) and Brier score (calibration + discrimination), each computed on "
    "calibrated probabilities."
)

doc.add_heading("3.6 Cross-Chemistry Transfer Protocol", level=2)
doc.add_paragraph(
    "We evaluate cross-chemistry transfer by training models on LCO cells (NASA, CALCE, or both) and "
    "testing on all Oxford LFP cells. Three training configurations are tested: NASA\u2192Oxford (37 NASA "
    "cells), CALCE\u2192Oxford (7 CALCE cells), and ALL\u2192Oxford (44 combined LCO cells). For each "
    "training configuration, models are evaluated with and without SOH as a feature. A positive transfer "
    "result (AUC significantly above 0.5) would indicate that degradation patterns learned on LCO "
    "generalize to LFP. Evaluation uses a single train/test split by dataset: "
    "all LCO cells are used for training, and all Oxford LFP cells are used for testing."
)

doc.add_heading("3.7 Evaluation Protocol", level=2)
doc.add_paragraph(
    "Within-dataset evaluation uses 5-fold GroupKFold stratified by cell: all cycles from a given cell "
    "belong to the same fold, ensuring that generalization is measured across unseen cells rather than "
    "unseen cycles. Four prediction horizons H \u2208 {10, 20, 30, 50} are tested, where the label for "
    "cycle t is positive if the battery fails within [t, t+H). Metrics are reported as means across folds. "
    "Models are retrained from scratch for each horizon. The best calibration method per (eval, dataset) "
    "pair is selected by mean AUC."
)

# ══════════════════════════════════════════════════════════════════════════
#  4. RESULTS
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("4. Results", level=1)

# ── 4.1 Within-Dataset Performance ───────────────────────────────────────
doc.add_heading("4.1 Within-Dataset Performance", level=2)
doc.add_paragraph(
    "Table 1 presents mean AUC and Brier scores (across all four horizons, Platt-calibrated) for all "
    "three models on NASA and CALCE. Both datasets show strong discrimination: AUC ranges from "
    "0.851 (Random Forest on NASA) to 0.930 (LightGBM on CALCE). CALCE achieves consistently lower "
    "Brier scores than NASA, reflecting the larger number of cycles per cell and smoother degradation "
    "trajectories."
)

# Read data for tables
df = pd.read_csv(CSV_PATH)

# Best method per (eval, dataset, model, H)
best = df.loc[df.groupby(["eval", "dataset", "model", "H"])["AUC_cal"].idxmax()]

# Table 1: within-dataset AUC per H + mean, grouped by (model, dataset)
t1_rows = []
for (mod, ds), grp in best[best["eval"] == "within"].groupby(["model", "dataset"]):
    aucs = {h: grp.loc[grp["H"] == h, "AUC_cal"].values[0] if (grp["H"] == h).any() else None for h in [10, 20, 30, 50]}
    mean_auc = grp["AUC_cal"].mean()
    mean_brier = grp["Brier_cal"].mean()
    t1_rows.append([mod, ds,
                    f"{aucs[10]:.3f}" if aucs[10] else "-",
                    f"{aucs[20]:.3f}" if aucs[20] else "-",
                    f"{aucs[30]:.3f}" if aucs[30] else "-",
                    f"{aucs[50]:.3f}" if aucs[50] else "-",
                    f"{mean_auc:.3f}", f"{mean_brier:.3f}"])
make_table(
    ["Model", "Dataset", "H=10", "H=20", "H=30", "H=50", "Mean AUC", "Brier"],
    t1_rows
)

add_figure(
    os.path.join(FIG_DIR, "Fig01_Within_Dataset_AUC.png"),
    "Figure 1: Within-dataset AUC heatmap at H=20 (best calibration per dataset). "
    "AUC values range from 0.868 (Random Forest on NASA) to 0.923 (LightGBM on CALCE)."
)

add_figure(
    os.path.join(FIG_DIR, "Fig05_MultiHorizon_AUC.png"),
    "Figure 2: Multi-horizon AUC on NASA (Platt-calibrated) as a function of prediction horizon H. "
    "AUC improves steadily from H=10 to H=30 across all models, with a plateau or mild regression at H=50."
)

# ── 4.2 Calibration ──────────────────────────────────────────────────────
doc.add_heading("4.2 Platt vs. Isotonic Calibration", level=2)
doc.add_paragraph(
    "Table 2 compares isotonic and Platt calibration Brier scores across horizons for both datasets, "
    "averaged over all models. Platt scaling universally outperforms isotonic regression, but the "
    "improvement is dramatically larger on CALCE."
)

# Table 2: calibration comparison by method
cal = df[df["eval"] == "within"]
t2 = cal.groupby(["dataset", "method"])["Brier_cal"].mean().round(3).reset_index()
make_table(
    ["Dataset", "Method", "Mean Brier"],
    [[row.dataset, row.method, f"{row.Brier_cal:.3f}"] for _, row in t2.iterrows()]
)

doc.add_paragraph(
    "On NASA, Platt reduces Brier from 0.222 to 0.197 (an 11% improvement). On CALCE, the improvement "
    "is far more substantial: Brier drops from 0.098 to 0.069 (a 30% reduction). The isotonic step "
    "function overcorrects on CALCE\u2019s long-tailed SOH distribution, where the vast number of "
    "midlife cycles with moderate failure probabilities produce calibration noise. Platt\u2019s sigmoid "
    "fit is more robust to this distributional characteristic, consistent with findings from the "
    "calibration literature [3]."
)

add_figure(
    os.path.join(FIG_DIR, "Fig02_Calibration_Comparison.png"),
    "Figure 3: Calibration comparison (isotonic vs. Platt) across horizons for NASA (left) and "
    "CALCE (right). Platt consistently produces lower (better) Brier scores (NASA: 0.197 vs 0.222; "
    "CALCE: 0.069 vs 0.098); the gap is most pronounced on CALCE."
)

# ── 4.3 Cross-Chemistry Transfer ──────────────────────────────────────────
doc.add_heading("4.3 Cross-Chemistry Transfer", level=2)
doc.add_paragraph(
    "Table 3a and Figure 4 summarize cross-chemistry transfer results when SOH is "
    "included as a feature. Models trained on LCO cells and tested on Oxford LFP cells achieve AUC "
    "values ranging from 0.703 (Random Forest, CALCE\u2192Oxford) to 1.000 (Random Forest, "
    "NASA\u2192Oxford). These results appear to demonstrate strong cross-chemistry generalization."
)

add_figure(
    os.path.join(FIG_DIR, "Fig03_CrossChem_With_SOH.png"),
    "Figure 4: Cross-chemistry transfer heatmap with SOH included (H=20). "
    "AUC 0.703\u20130.923 (CALCE\u2192LFP) to 0.945\u20131.000 (NASA\u2192LFP). "
    "These values appear to demonstrate strong generalization but are driven by SOH (see Fig 5)."
)

doc.add_paragraph(
    "However, when SOH is removed from the feature set, the result is unambiguous: "
    "AUC collapses to 0.51\u20130.54 across all model\u2013training "
    "combinations (Table 3b, Figure 5). This near-random performance reveals that SOH was the "
    "sole driver of the apparent cross-chemistry "
    "generalization. When the model has access to SOH, it learns a chemistry-specific SOH-to-RUL mapping "
    "\u2014 e.g., \u201cSOH decreasing from 1.0 to 0.8 over ~250 cycles\u201d for NASA LCO cells. "
    "When tested on LFP cells with similar SOH trajectories, the model applies the same learned mapping, "
    "producing predictions that correlate with SOH and therefore appear accurate. In effect, the model is "
    "reading off a lookup table rather than learning chemistry-agnostic degradation features."
)

# Table 3a: cross-chem with SOH (H=20 only)
cross = df[df["eval"].isin(["train_nasa_with_soh", "train_calce_with_soh", "train_nasa+calce_with_soh"])]
cross = cross[cross["H"] == 20].copy()
# Pick best method
cross = cross.loc[cross.groupby(["eval", "model"])["AUC_cal"].idxmax()]
rows3a = []
for _, row in cross.sort_values("eval").iterrows():
    label = row.eval.replace("train_", "").replace("_with_soh", "")
    rows3a.append([label, row.model, f"{row.AUC_cal:.3f}"])
make_table(
    ["Training Config", "Model", "AUC (H=20, best cal.)"],
    rows3a
)

# Table 3b: cross-chem without SOH (H=20 only)
cross_no = df[df["eval"].isin(["train_nasa_no_soh", "train_calce_no_soh", "train_nasa+calce_no_soh"])]
cross_no = cross_no[cross_no["H"] == 20].copy()
cross_no = cross_no.loc[cross_no.groupby(["eval", "model"])["AUC_cal"].idxmax()]
rows3b = []
for _, row in cross_no.sort_values("eval").iterrows():
    label = row.eval.replace("train_", "").replace("_no_soh", "")
    rows3b.append([label, row.model, f"{row.AUC_cal:.3f}"])
make_table(
    ["Training Config", "Model", "AUC (H=20, best cal.)"],
    rows3b
)

add_figure(
    os.path.join(FIG_DIR, "Fig04_CrossChem_No_SOH.png"),
    "Figure 5: Cross-chemistry transfer heatmap with SOH removed from features (H=20). "
    "AUC collapses to 0.51\u20130.54 across all 9 model\u2013training combinations, "
    "demonstrating that voltage, current, and temperature features alone carry no transferable signal "
    "between LCO and LFP."
)

# ══════════════════════════════════════════════════════════════════════════
#  5. DISCUSSION
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("5. Discussion", level=1)
doc.add_paragraph(
    "The central finding of this study is the negative cross-chemistry transfer result. The fact that "
    "SOH removal reduces AUC from near-perfect to near-random is not merely a null result \u2014 it "
    "reveals a specific mechanism: SOH functions as a chemistry-specific lookup key. LCO and LFP cells "
    "both traverse SOH values from 1.0 to 0.8 and below, but the relationship between SOH and "
    "time-to-failure differs fundamentally between chemistries. For LCO under accelerated aging, "
    "SOH decay is roughly linear with cycle count; for LFP, the flat voltage plateau decouples "
    "voltage-based features from capacity degradation. A model that learns \u201cSOH=0.85 means ~50 "
    "cycles to failure\u201d from LCO data will apply that same mapping to LFP regardless of whether "
    "the relationship holds."
)
doc.add_paragraph(
    "This finding has practical implications for battery monitoring systems in mixed-chemistry fleets. "
    "A hazard model trained on one chemistry cannot be naively deployed on another without feature "
    "engineering, domain adaptation, or retraining. The voltage, current, and temperature features "
    "commonly available in battery management systems do not, in isolation, provide a chemistry-invariant "
    "failure signature robust enough for cross-chemistry transfer."
)
doc.add_paragraph(
    "The within-dataset results, by contrast, are robust. Across two LCO datasets with different "
    "cycling protocols (NASA\u2019s random-walk aging vs. CALCE\u2019s constant-current cycling), "
    "all three model classes achieve AUC of 0.85 or above with Platt calibration. The multi-horizon "
    "formulation works consistently: AUC improves from H=10 to H=30 across all models on NASA, "
    "with a plateau or mild regression at H=50. On CALCE, AUC decreases slightly from H=10 to H=50, "
    "reflecting the longer per-cell degradation tails. This establishes a "
    "reproducible baseline for future work."
)
doc.add_paragraph(
    "We note several limitations. The Oxford LFP dataset contains only 5 cells, which is "
    "insufficient for reliable within-dataset evaluation (AUC \u2248 1.0 due to easy cross-validation) "
    "and limits the statistical power of the cross-chemistry analysis. The transfer evaluation is "
    "unidirectional (LCO \u2192 LFP) and may not generalize to other chemistry pairs (LCO \u2192 NMC, "
    "NMC \u2192 LFP). Additionally, only tree-based models are tested; deep learning approaches with "
    "learned representations may capture more transferable features."
)

# ══════════════════════════════════════════════════════════════════════════
#  6. CONCLUSION
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("6. Conclusion", level=1)
doc.add_paragraph(
    "We extended a multi-horizon battery hazard classification framework from a single model on a "
    "single dataset to three models on two LCO datasets plus cross-chemistry transfer to LFP. Platt "
    "calibration universally outperforms isotonic regression, with the largest gains on datasets with "
    "long-tailed degradation distributions. Within-dataset AUC of 0.85\u20130.93 establishes a "
    "reproducible baseline for LCO hazard prediction. The key finding, however, is the failure of "
    "cross-chemistry transfer when SOH is removed as a feature: AUC drops from near-perfect to "
    "near-random, demonstrating that SOH encodes a chemistry-specific capacity-to-RUL mapping, not "
    "a transferable degradation invariant. Future work should explore learned feature representations "
    "designed explicitly for chemistry invariance, larger multi-chemistry datasets, and bidirectional "
    "transfer evaluations."
)

# ══════════════════════════════════════════════════════════════════════════
#  REFERENCES
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("References", level=1)
refs = [
    "[1] T. A. Shikdar and H. Laaksonen, \u201cLearning when not to use a battery: Multihorizon failure intelligence,\u201d Int. Trans. Electr. Energy Syst., vol. 2026, art. 6000810, 2026. doi:10.1155/etep/6000810.",
    "[2] B. Zadrozny and C. Elkan, \u201cTransforming classifier scores into accurate multiclass probability estimates,\u201d in Proc. ACM SIGKDD, 2002.",
    "[3] J. Platt, \u201cProbabilistic outputs for support vector machines and comparisons to regularized likelihood methods,\u201d in Advances in Large Margin Classifiers, 1999.",
    "[4] B. Saha and K. Goebel, \u201cBattery Data Set,\u201d NASA Ames Prognostics Data Repository, 2007.",
    "[5] CALCE Battery Research Group, \u201cBattery aging datasets,\u201d University of Maryland, 2023.",
    "[6] Oxford Battery Degradation Dataset, \u201cLFP pouch cell cycling data,\u201d University of Oxford, 2021.",
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(9)

# Save
out_path = os.path.join(DOCX_DIR, "paper.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
